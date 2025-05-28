import base64
from enum import Enum
from io import BytesIO

from docx import Document
from PIL import Image


class EChatModel(Enum):
    AZURE = 'AZURE'
    AWS = 'AWS'


async def extract_and_upload_images(doc, user_id: str, chat_model: EChatModel):
    from services import calculate_image_token, upload_file_to_s3
    from services.preprocess import resize_image
    token_image = 0
    data = []

    for index, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_content_type = rel.target_part.content_type
            image_stream = BytesIO(image_data)

            image = Image.open(image_stream)
            new_width, new_height, token = calculate_image_token(
                image.width, image.height, chat_model
            )
            token_image += token

            resized_img = resize_image(image, new_width, new_height)
            
            resized_stream = BytesIO()
            resized_img.save(resized_stream, format=image.format)
            resized_stream.seek(0)
            
            if chat_model == EChatModel.AZURE:
                image_name = f"{index}_{rel.target_ref.split('/')[-1]}"
                file_url = await upload_file_to_s3(resized_stream, image_name, user_id)
                data.append(file_url)
            elif chat_model == EChatModel.AWS:
                base64_image = base64.b64encode(resized_stream.read()).decode('utf-8')
                data_uri = f"data:{image_content_type};base64,{base64_image}"
                data.append(data_uri)

    return data, token_image

async def extract_text_from_docx(doc_path: BytesIO, user_id: str, chat_model: EChatModel):
    doc = Document(doc_path)
    data = []
    
    image_urls, token_image = await extract_and_upload_images(doc, user_id, chat_model)
    
    paragraph_group = []
    
    data = await process_document_elements(doc, paragraph_group, image_urls)
    
    return data, token_image

async def process_document_elements(doc, paragraph_group: list, image_urls: list):
    data = []
    image_index = 0
    table_num = 0
    para_num = 0
    
    for idx, element in enumerate(doc.element.body):
        tag_name = element.tag.split("}")[-1]
        
        if idx >= len(doc.paragraphs):
            continue
            
        if 'graphicData' in doc.paragraphs[idx]._p.xml:
            data = append_text_and_image(data, paragraph_group, image_urls, image_index)
            paragraph_group.clear()
            image_index += 1
            
        if tag_name == "p":
            process_paragraph(doc, para_num, paragraph_group)
            para_num += 1
        elif tag_name == "tbl":
            process_table(doc, table_num, paragraph_group)
            table_num += 1
    
    if paragraph_group:
        data.append({
            "type": "text",
            "text": "\n".join(paragraph_group)
        })
    
    return data

def append_text_and_image(data, paragraph_group: list, image_urls: list, image_index: int) -> list:
    if paragraph_group:
        data.append({
            "type": "text",
            "text": "\n".join(paragraph_group)
        })
    
    if image_index < len(image_urls):
        data.append({
            "type": "image_url",
            "image_url": image_urls[image_index]
        })
    
    return data

def process_paragraph(doc, para_num: int, paragraph_group: list[str]) -> None:
    para = doc.paragraphs[para_num]
    text = para.text.strip()
    
    if not text:
        return
        
    style = get_paragraph_style(para)
    formatted_text = format_paragraph_by_style(text, style)
    paragraph_group.append(formatted_text)

def get_paragraph_style(para) -> str:
    try:
        return para.style.name if hasattr(para.style, 'name') else "Normal"
    except AttributeError:
        return "Normal"

def format_paragraph_by_style(text: str, style: str) -> str:
    if style.startswith("Heading"):
        level = int(style.replace("Heading ", "")) if style.replace("Heading ", "").isdigit() else 2
        return f"{'#' * level}> {text}"
    elif style == "Title":
        return f"<<{text}>>\n"
    elif style == "List Paragraph":
        return f"  - {text}"
    else:
        return f"{text}\n"

def process_table(doc, table_num: int, paragraph_group: list[str]) -> None:
    table = doc.tables[table_num]
    md_content = convert_table_to_markdown(table)
    paragraph_group.append(md_content)

def convert_table_to_markdown(table) -> str:
    md_content = ""
    rows = table.rows
    
    if not rows:
        return md_content
        
    headers = [cell.text.strip() for cell in rows[0].cells]
    md_content += "| " + " | ".join(headers) + " |\n"
    md_content += "|" + "|".join(["---"] * len(headers)) + "|\n"
    
    for row in rows[1:]:
        row_data = [cell.text.strip() for cell in row.cells]
        md_content += "| " + " | ".join(row_data) + " |\n"
        
    return md_content